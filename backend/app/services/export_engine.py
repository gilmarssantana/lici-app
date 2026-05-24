from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from fastapi import HTTPException

from app.schemas.case import CaseRecord
from app.schemas.document_generator import GeneratedDocumentRecord
from app.services.audit_log import audit_event
from app.services.case_store import JsonCaseStore
from app.services.document_generator_store import HybridGeneratedDocumentStore
from app.services.document_pg_store import PostgresExportedFileStore

EXPORT_DIR = Path("/root/lici-app/storage/exportados")


class LiciExportEngine:
    def __init__(
        self,
        document_store: HybridGeneratedDocumentStore | None = None,
        case_store: JsonCaseStore | None = None,
        export_dir: Path | str = EXPORT_DIR,
    ):
        self.document_store = document_store or HybridGeneratedDocumentStore()
        self.case_store = case_store or JsonCaseStore()
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)
        self.export_store = PostgresExportedFileStore()

    def engine_info(self) -> dict[str, object]:
        return {
            "nome": "LICI Export Engine",
            "objetivo": "Exportar peças e relatórios da LICI em TXT/DOCX para uso operacional real.",
            "endpoints": [
                "GET /export/engine",
                "GET /export/documentos/{id}/txt",
                "GET /export/documentos/{id}/docx",
                "GET /export/casos/{id}/relatorio-txt",
                "GET /export/casos/{id}/relatorio-docx",
            ],
            "formatos": ["txt", "docx"],
            "persistencia": str(self.export_dir),
            "integracoes": ["Document Generator", "Case Engine", "Audit Log"],
        }

    def export_document_txt(self, document_id: str) -> Path:
        doc = self._get_document(document_id)
        path = self._export_path(f"peca-{doc.tipo}-{doc.id}", "txt")
        path.write_text(doc.texto, encoding="utf-8")
        self._audit("export_documento_txt", "document_generator", doc.id, path, {"tipo": doc.tipo, "titulo": doc.titulo})
        return path

    def export_document_docx(self, document_id: str) -> Path:
        doc = self._get_document(document_id)
        path = self._export_path(f"peca-{doc.tipo}-{doc.id}", "docx")
        self._write_docx(path, doc.titulo, doc.texto)
        self._audit("export_documento_docx", "document_generator", doc.id, path, {"tipo": doc.tipo, "titulo": doc.titulo})
        return path

    def export_case_report_txt(self, case_id: str) -> Path:
        case = self._get_case(case_id)
        texto = self._case_report(case)
        path = self._export_path(f"relatorio-caso-{case.id}", "txt")
        path.write_text(texto, encoding="utf-8")
        self._audit("export_relatorio_caso_txt", "case_engine", case.id, path, {"orgao": case.orgao, "objeto": case.objeto})
        return path

    def export_case_report_docx(self, case_id: str) -> Path:
        case = self._get_case(case_id)
        texto = self._case_report(case)
        path = self._export_path(f"relatorio-caso-{case.id}", "docx")
        self._write_docx(path, f"Relatório do Caso — {case.orgao}", texto)
        self._audit("export_relatorio_caso_docx", "case_engine", case.id, path, {"orgao": case.orgao, "objeto": case.objeto})
        return path

    def _get_document(self, document_id: str) -> GeneratedDocumentRecord:
        for doc in self.document_store.list():
            if doc.id == document_id:
                return doc
        raise HTTPException(status_code=404, detail="documento gerado não encontrado")

    def _get_case(self, case_id: str) -> CaseRecord:
        case = self.case_store.get(case_id)
        if case is None:
            raise HTTPException(status_code=404, detail="caso vivo não encontrado")
        return case

    def _case_report(self, case: CaseRecord) -> str:
        lines = [
            "RELATÓRIO OPERACIONAL — CASO VIVO LICI",
            "",
            f"ID: {case.id}",
            f"Cliente: {case.cliente}",
            f"Órgão: {case.orgao}",
            f"Objeto: {case.objeto}",
            f"Modalidade: {case.modalidade or 'não informada'}",
            f"Status: {case.status}",
            f"Fase atual: {case.fase_atual}",
            f"Score estratégico: {case.score_estrategico}/100",
            f"Criado em: {case.criado_em}",
            f"Atualizado em: {case.atualizado_em}",
            "",
            "RISCOS",
            self._bullets(case.riscos),
            "",
            "OPORTUNIDADES",
            self._bullets(case.oportunidades),
            "",
            "MEMÓRIAS RELACIONADAS",
            self._bullets([str(item.get('titulo') or item.get('id') or item) for item in case.memorias_relacionadas]),
            "",
            "HISTÓRICO OPERACIONAL",
        ]
        if case.historico_operacional:
            for event in case.historico_operacional:
                lines.extend([
                    f"- {event.data} | {event.tipo} | fase: {event.fase}",
                    f"  Descrição: {event.descricao}",
                    f"  Impacto: {event.impacto or 'não informado'}",
                    f"  Aprendizado: {event.aprendizado_operacional or 'não informado'}",
                ])
        else:
            lines.append("- Sem eventos registrados.")
        if case.memoria_sugerida:
            lines.extend([
                "",
                "MEMÓRIA SUGERIDA",
                f"Tipo: {case.memoria_sugerida.tipo}",
                f"Título: {case.memoria_sugerida.titulo}",
                f"Descrição: {case.memoria_sugerida.descricao}",
                f"Estratégia: {case.memoria_sugerida.estrategia}",
                f"Uso futuro: {case.memoria_sugerida.uso_futuro}",
            ])
        return "\n".join(lines).strip() + "\n"

    def _write_docx(self, path: Path, title: str, text: str) -> None:
        document = Document()
        document.add_heading(title or "Documento LICI", level=1)
        for block in text.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if len(block) <= 90 and block.upper() == block and not block.startswith("-"):
                document.add_heading(block.title(), level=2)
                continue
            for line in block.splitlines():
                if line.startswith("- "):
                    document.add_paragraph(line[2:], style="List Bullet")
                else:
                    document.add_paragraph(line)
        document.save(path)

    def _export_path(self, base: str, ext: str) -> Path:
        stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S")
        safe = self._safe_slug(base)
        return self.export_dir / f"{safe}-{stamp}.{ext}"

    def _safe_slug(self, value: str) -> str:
        return re.sub(r"[^A-Za-z0-9À-ÿ._-]+", "-", value).strip("-.").lower()[:150] or "export-lici"

    def _bullets(self, items: list[str]) -> str:
        return "\n".join(f"- {item}" for item in items if item) or "- Não registrado."

    def _audit(self, acao: str, modulo_origem: str, related_id: str, path: Path, detalhes: dict) -> None:
        payload = {"arquivo": str(path), "modulo_origem": modulo_origem, **detalhes}
        audit_event("export_engine", acao, "ok", payload, related_id)
        try:
            if self.export_store.available():
                self.export_store.create(
                    file_path=str(path),
                    file_type=path.suffix.lstrip(".").lower(),
                    source_module=modulo_origem,
                    source_id=related_id,
                    title=detalhes.get("titulo") or detalhes.get("orgao") or path.name,
                    metadata={"acao": acao, **detalhes},
                )
                audit_event("export_engine", "dual_write_postgres", "ok", {"operacao": acao, "arquivo": str(path)}, related_id)
            else:
                audit_event("export_engine", "dual_write_postgres", "erro", {"operacao": acao, "motivo": "postgres_indisponivel", "arquivo": str(path)}, related_id)
        except Exception as exc:
            audit_event("export_engine", "dual_write_postgres", "erro", {"operacao": acao, "erro": str(exc), "arquivo": str(path)}, related_id)
